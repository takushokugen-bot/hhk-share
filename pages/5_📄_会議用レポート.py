import streamlit as st
from supabase import create_client, Client
import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO
import datetime as dt

from modules.font import *   # ← フォント設定は必ず最後に import
from modules.ai_analysis import summarize_reports   # ← AI要約

# ============================
# Supabase 接続
# ============================

SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

st.title("📄 会議用レポート出力（Word）")

# ============================
# 期間指定
# ============================

today = dt.date.today()
start_date = st.date_input("開始日", today.replace(day=1))
end_date = st.date_input("終了日", today)

# ============================
# データ取得
# ============================

reports = (
    supabase.table("hhk_reports")
    .select("id, reported_at, description, reporter, company, locations(name), categories(name)")
    .order("reported_at")
    .execute()
    .data
)

if not reports:
    st.info("まだ投稿がありません。")
    st.stop()

# ============================
# DataFrame 化
# ============================

df = pd.DataFrame([
    {
        "ID": r["id"],
        "発生日時": pd.to_datetime(r["reported_at"]),
        "場所": r["locations"]["name"] if r["locations"] else "（未設定）",
        "カテゴリ": r["categories"]["name"] if r["categories"] else "（未設定）",
        "内容": r["description"] if r["description"] else "（未入力）",
        "投稿者": r["reporter"] if r["reporter"] else "（未入力）",
        "会社名": r["company"] if r["company"] else "（未入力）",
    }
    for r in reports
])

# 期間で絞る
df = df[
    (df["発生日時"].dt.date >= start_date) &
    (df["発生日時"].dt.date <= end_date)
]

if df.empty:
    st.warning("この期間のデータはありません。")
    st.stop()

# ============================
# 月次KPI（件数一覧＋増減率）
# ============================

df["月"] = df["発生日時"].dt.to_period("M").astype(str)

monthly_counts = df.groupby("月")["ID"].count().reset_index()
monthly_counts = monthly_counts.sort_values("月")

monthly_counts["前月比"] = monthly_counts["ID"].pct_change() * 100
monthly_counts["前月比"] = monthly_counts["前月比"].apply(
    lambda x: f"{x:.1f}%" if pd.notnull(x) else "—"
)

# 月次推移グラフ
fig_monthly = plt.figure(figsize=(8, 4))
plt.plot(monthly_counts["月"], monthly_counts["ID"], marker="o")
plt.title("月次件数推移")
plt.xlabel("月")
plt.ylabel("件数")
plt.grid(True)

# ============================
# グラフ生成
# ============================

def plot_bar(data, x, title):
    fig, ax = plt.subplots(figsize=(8, 4))
    sns.countplot(data=data, x=x, ax=ax)
    ax.set_title(title)
    plt.xticks(rotation=90)
    return fig

def plot_heatmap(pivot, title):
    fig, ax = plt.subplots(figsize=(10, 5))
    sns.heatmap(pivot, cmap="Reds", annot=True, fmt="d", ax=ax)
    ax.set_title(title)
    return fig

def fig_to_png_bytes(fig):
    buf = BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    buf.seek(0)
    return buf

df["時間帯"] = df["発生日時"].dt.hour
df["曜日"] = df["発生日時"].dt.day_name()

fig_time = plot_bar(df, "時間帯", "時間帯別件数")
fig_week = plot_bar(df, "曜日", "曜日別件数")
fig_loc = plot_bar(df, "場所", "場所別件数")
fig_cat = plot_bar(df, "カテゴリ", "カテゴリ別件数")
fig_company = plot_bar(df, "会社名", "会社別件数")

pivot_week = df.pivot_table(index="曜日", columns="時間帯", aggfunc="size", fill_value=0)
pivot_loc = df.pivot_table(index="場所", columns="時間帯", aggfunc="size", fill_value=0)

fig_heat_week = plot_heatmap(pivot_week, "曜日 × 時間帯 ヒートマップ")
fig_heat_loc = plot_heatmap(pivot_loc, "場所 × 時間帯 ヒートマップ")

# ============================
# 🤖 AI要約（期間全体）
# ============================

with st.spinner("AIが期間全体の要約を作成中..."):
    ai_summary_all = summarize_reports(df)

# ============================
# 🤖 月別AI要約
# ============================

monthly_ai = {}
for month in monthly_counts["月"].unique():
    df_month = df[df["月"] == month]
    with st.spinner(f"AIが {month} の要約を作成中..."):
        monthly_ai[month] = summarize_reports(df_month)

# ============================
# Word レポート生成
# ============================

def build_doc(df_export):
    doc = Document()

    doc.add_heading("HHK（ヒヤリハット）振り返りレポート", level=1)
    doc.add_paragraph(f"期間：{start_date} ～ {end_date}")
    doc.add_paragraph("")

    # 0. AI要約（期間全体）
    doc.add_heading("0. AI要約（期間全体）", level=2)
    doc.add_paragraph(ai_summary_all)
    doc.add_paragraph("")

    # 1. KPI
    doc.add_heading("1. KPI（重要指標：月次比較）", level=2)
    doc.add_paragraph("■ 月ごとの件数・前月比")
    for _, row in monthly_counts.iterrows():
        doc.add_paragraph(
            f"- {row['月']}：{row['ID']} 件（前月比：{row['前月比']}）",
            style="List Bullet"
        )

    doc.add_paragraph("■ 月次件数推移グラフ")
    doc.add_picture(fig_to_png_bytes(fig_monthly), width=Inches(6))
    doc.add_paragraph("")

    # 2. 集計サマリ
    doc.add_heading("2. 集計サマリ", level=2)
    doc.add_paragraph(f"総件数：{len(df_export)} 件")

    by_loc = df_export.groupby("場所")["ID"].count().reset_index().sort_values("ID", ascending=False)
    doc.add_paragraph("■ 場所別件数")
    for _, row in by_loc.iterrows():
        doc.add_paragraph(f"- {row['場所']}：{row['ID']} 件", style="List Bullet")
    doc.add_paragraph("")

    by_cat = df_export.groupby("カテゴリ")["ID"].count().reset_index().sort_values("ID", ascending=False)
    doc.add_paragraph("■ カテゴリ別件数")
    for _, row in by_cat.iterrows():
        doc.add_paragraph(f"- {row['カテゴリ']}：{row['ID']} 件", style="List Bullet")
    doc.add_paragraph("")

    by_company = df_export.groupby("会社名")["ID"].count().reset_index().sort_values("ID", ascending=False)
    doc.add_paragraph("■ 会社別件数")
    for _, row in by_company.iterrows():
        doc.add_paragraph(f"- {row['会社名']}：{row['ID']} 件", style="List Bullet")
    doc.add_paragraph("")

    # 3. グラフ
    doc.add_heading("3. グラフ", level=2)
    for title, fig in [
        ("時間帯別件数", fig_time),
        ("曜日別件数", fig_week),
        ("場所別件数", fig_loc),
        ("カテゴリ別件数", fig_cat),
        ("会社別件数", fig_company),
    ]:
        doc.add_paragraph(f"■ {title}")
        doc.add_picture(fig_to_png_bytes(fig), width=Inches(6))

    # 4. ヒートマップ
    doc.add_heading("4. ヒートマップ", level=2)
    doc.add_paragraph("■ 曜日 × 時間帯")
    doc.add_picture(fig_to_png_bytes(fig_heat_week), width=Inches(6))

    doc.add_paragraph("■ 場所 × 時間帯")
    doc.add_picture(fig_to_png_bytes(fig_heat_loc), width=Inches(6))

    # 5. 個別事例一覧
    doc.add_heading("5. 個別事例一覧", level=2)
    for _, row in df_export.sort_values("発生日時").iterrows():
        doc.add_paragraph(f"[ID {row['ID']}] {row['発生日時']}")
        doc.add_paragraph(f"場所：{row['場所']} / カテゴリ：{row['カテゴリ']}")
        doc.add_paragraph(f"会社：{row['会社名']}")
        doc.add_paragraph(f"内容：{row['内容']}")
        doc.add_paragraph(f"投稿者：{row['投稿者']}")
        doc.add_paragraph("-" * 40)

    # 6. 月別サマリ・月別AI要約・月別グラフ
    doc.add_heading("6. 月別サマリ・AI要約・月別グラフ", level=2)

    for month in monthly_counts["月"].unique():
        doc.add_heading(f"■ {month} のAI要約", level=3)
        doc.add_paragraph(monthly_ai[month])
        doc.add_paragraph("")

        doc.add_heading(f"■ {month} のサマリ", level=3)

        df_month = df_export[df_export["月"] == month]

        doc.add_paragraph(f"- 件数：{len(df_month)} 件")

        by_loc_m = df_month.groupby("場所")["ID"].count().reset_index().sort_values("ID", ascending=False)
        doc.add_paragraph("・場所別件数")
        for _, row in by_loc_m.iterrows():
            doc.add_paragraph(f"  - {row['場所']}：{row['ID']} 件", style="List Bullet")

        by_cat_m = df_month.groupby("カテゴリ")["ID"].count().reset_index().sort_values("ID", ascending=False)
        doc.add_paragraph("・カテゴリ別件数")
        for _, row in by_cat_m.iterrows():
            doc.add_paragraph(f"  - {row['カテゴリ']}：{row['ID']} 件", style="List Bullet")

        by_company_m = df_month.groupby("会社名")["ID"].count().reset_index().sort_values("ID", ascending=False)
        doc.add_paragraph("・会社別件数")
        for _, row in by_company_m.iterrows():
            doc.add_paragraph(f"  - {row['会社名']}：{row['ID']} 件", style="List Bullet")

        doc.add_paragraph("")

        doc.add_heading(f"■ {month} のグラフ", level=3)

        fig_time_m = plot_bar(df_month, "時間帯", f"{month} 時間帯別件数")
        fig_week_m = plot_bar(df_month, "曜日", f"{month} 曜日別件数")
        fig_loc_m = plot_bar(df_month, "場所", f"{month} 場所別件数")
        fig_cat_m = plot_bar(df_month, "カテゴリ", f"{month} カテゴリ別件数")
        fig_company_m = plot_bar(df_month, "会社名", f"{month} 会社別件数")

        for title, fig in [
            (f"{month} 時間帯別件数", fig_time_m),
            (f"{month} 曜日別件数", fig_week_m),
            (f"{month} 場所別件数", fig_loc_m),
            (f"{month} カテゴリ別件数", fig_cat_m),
            (f"{month} 会社別件数", fig_company_m),
        ]:
            doc.add_paragraph(f"● {title}")
            doc.add_picture(fig_to_png_bytes(fig), width=Inches(6))

        doc.add_paragraph("")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ============================
# ダウンロード
# ============================

doc_bytes = build_doc(df)
st.download_button(
    label="📥 Wordレポートをダウンロード",
    data=doc_bytes,
    file_name=f"HHKレポート_{start_date}_{end_date}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
